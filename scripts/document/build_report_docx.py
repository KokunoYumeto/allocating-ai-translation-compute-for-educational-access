#!/usr/bin/env python3
"""Build the educational-access research report as an accessible, styled DOCX.

The source of truth remains PAPER.md.  This builder deliberately implements a
small, deterministic subset of Markdown used by that manuscript: headings,
paragraphs, emphasis, inline code, ordered/unordered lists, pipe tables,
display equations, images, and figure/table captions.

Design authority:
  - documents skill preset: narrative_proposal
  - first-page pattern: editorial_cover
  - named override: landscape_wide_table (US Letter landscape, 9 in usable)

The script contains no publication logic and changes no source files.
"""

from __future__ import annotations

import argparse
import json
import math
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


PRESET = {
    "name": "narrative_proposal",
    "page": {
        "portrait_width_in": 8.5,
        "portrait_height_in": 11.0,
        "landscape_width_in": 11.0,
        "landscape_height_in": 8.5,
        "margin_in": 1.0,
        "header_footer_in": 0.492,
        "portrait_width_dxa": 9360,
        "landscape_width_dxa": 12960,
    },
    "body": {
        "font": "Calibri",
        "size_pt": 11,
        "before_pt": 0,
        "after_pt": 8,
        "line": 1.333,
        "alignment": "justified",
    },
    "headings": {
        "h1": {"size": 16, "color": "2E74B5", "before": 18, "after": 10},
        "h2": {"size": 13, "color": "2E74B5", "before": 12, "after": 6},
        "h3": {"size": 12, "color": "1F4D78", "before": 8, "after": 4},
    },
    "lists": {
        "marker_dxa": 261,
        "text_dxa": 540,
        "hanging_dxa": 279,
        "after_pt": 4,
        "line": 1.208,
    },
    "tables": {
        "indent_dxa": 120,
        "cell_top_bottom_dxa": 80,
        "cell_start_end_dxa": 120,
        "header_fill": "F4F6F9",
        "border": "C8D2DF",
    },
    "colors": {
        "ink": "203748",
        "heading": "2E74B5",
        "dark_heading": "1F4D78",
        "muted": "66717E",
        "gold": "8A6A12",
        "abstract_border": "7FA6C8",
        "abstract_fill": "F4F7FA",
    },
}

ALT_TEXT_DEFAULT = (
    "Horizontal interval chart for the headline Top 20 interventions. Each line "
    "shows the range of base, optimistic, and scarcity ranks; a colored point "
    "marks the admission-lane rank. The all-zero conservative rank is not plotted."
)


@dataclass
class Token:
    kind: str
    value: object
    level: int = 0


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, color: str = "000000") -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 12) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = pbdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        pbdr.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "5")
    edge.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, PRESET["body"]["font"], PRESET["body"]["size_pt"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Calibri", 30, PRESET["colors"]["ink"])
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(9)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 15, "2B5163")
    subtitle.font.italic = False
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(26)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, key in enumerate(("h1", "h2", "h3"), start=1):
        spec = PRESET["headings"][key]
        style = styles[f"Heading {index}"]
        set_style_font(style, "Calibri", spec["size"], spec["color"])
        style.font.bold = True
        style.paragraph_format.space_before = Pt(spec["before"])
        style.paragraph_format.space_after = Pt(spec["after"])
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True
        # Remove template residue explicitly. Page starts are controlled by the
        # builder's section/page-break plan, never by inherited Heading styles.
        style.paragraph_format.page_break_before = False

    caption = styles["Caption"]
    set_style_font(caption, "Calibri", 9, PRESET["colors"]["muted"])
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.widow_control = True

    if "Proposal List" not in styles:
        list_style = styles.add_style("Proposal List", WD_STYLE_TYPE.PARAGRAPH)
    else:
        list_style = styles["Proposal List"]
    set_style_font(list_style, "Calibri", 11)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(PRESET["lists"]["after_pt"])
    list_style.paragraph_format.line_spacing = PRESET["lists"]["line"]
    list_style.paragraph_format.widow_control = True

    if "Reference Entry" not in styles:
        ref_style = styles.add_style("Reference Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref_style = styles["Reference Entry"]
    set_style_font(ref_style, "Calibri", 9.5)
    ref_style.paragraph_format.left_indent = Inches(0.25)
    ref_style.paragraph_format.first_line_indent = Inches(-0.25)
    ref_style.paragraph_format.space_before = Pt(0)
    ref_style.paragraph_format.space_after = Pt(6)
    ref_style.paragraph_format.line_spacing = 1.10
    ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_style.paragraph_format.widow_control = True

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    set_style_font(equation, "Cambria Math", 10.5, PRESET["colors"]["ink"])
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(7)
    equation.paragraph_format.space_after = Pt(9)
    equation.paragraph_format.keep_together = True

    if "Table Source Note" not in styles:
        note = styles.add_style("Table Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Table Source Note"]
    set_style_font(note, "Calibri", 9, PRESET["colors"]["muted"])
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.0


def _max_numbering_id(numbering, element_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(f"w:{element_name}")):
        raw = element.get(qn(f"w:{attr_name}"))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0)


class NumberingFactory:
    def __init__(self, doc: Document) -> None:
        self.numbering = doc.part.numbering_part.element
        next_abs = _max_numbering_id(self.numbering, "abstractNum", "abstractNumId") + 1
        self.abstract = {
            False: self._add_abstract(next_abs, ordered=False),
            True: self._add_abstract(next_abs + 1, ordered=True),
        }

    def _add_abstract(self, abstract_id: int, ordered: bool) -> int:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "%1." if ordered else "•")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(PRESET["lists"]["text_dxa"]))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(PRESET["lists"]["text_dxa"]))
        ind.set(qn("w:hanging"), str(PRESET["lists"]["hanging_dxa"]))
        ppr.append(ind)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rpr.append(rfonts)
        lvl.append(rpr)
        abstract.append(lvl)
        self.numbering.append(abstract)
        return abstract_id

    def new_instance(self, ordered: bool) -> int:
        num_id = _max_numbering_id(self.numbering, "num", "numId") + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(self.abstract[ordered]))
        num.append(abstract_id)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
        self.numbering.append(num)
        return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells = re.split(r"(?<!\\)\|", line)
    return [cell.strip().replace(r"\|", "|") for cell in cells]


def is_table_separator(cells: Sequence[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def parse_markdown(text: str) -> tuple[dict[str, str], list[Token]]:
    lines = text.replace("\r\n", "\n").split("\n")
    metadata = {"title": "Research report", "subtitle": "", "date": ""}
    if lines and lines[0].startswith("# "):
        metadata["title"] = lines[0][2:].strip()
    for line in lines[1:12]:
        if line.startswith("## ") and line[3:].strip().lower() != "abstract":
            metadata["subtitle"] = line[3:].strip()
        date_match = re.search(r"\*\*Date:\*\*\s*(.+)", line)
        if date_match:
            metadata["date"] = date_match.group(1).strip().rstrip("  ")

    start = next((i for i, line in enumerate(lines) if line.strip() == "## Abstract"), 0)
    tokens: list[Token] = []
    i = start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            raw_level = len(heading.group(1))
            # The manuscript reserves # for the cover title and appendices; body ##
            # is the top navigational level in the produced report.
            level = 1 if raw_level <= 2 else min(3, raw_level - 1)
            tokens.append(Token("heading", heading.group(2).strip(), level))
            i += 1
            continue

        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = split_table_row(lines[i])
                if not is_table_separator(cells):
                    rows.append(cells)
                i += 1
            if rows:
                width = max(len(row) for row in rows)
                rows = [row + [""] * (width - len(row)) for row in rows]
                tokens.append(Token("table", rows))
            continue

        if stripped == r"\[":
            equation_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                equation_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                i += 1
            tokens.append(Token("equation", " ".join(equation_lines).strip()))
            continue

        if stripped.startswith("!["):
            image_lines = [stripped]
            while i + 1 < len(lines) and "](" not in " ".join(image_lines):
                i += 1
                image_lines.append(lines[i].strip())
            joined = " ".join(image_lines)
            match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", joined)
            if match:
                tokens.append(Token("image", {"alt": match.group(1), "path": match.group(2)}))
            else:
                tokens.append(Token("paragraph", joined))
            i += 1
            continue

        block: list[str] = []
        while i < len(lines):
            candidate = lines[i]
            c = candidate.strip()
            if not c:
                break
            if block and (
                re.match(r"^#{1,6}\s+", c)
                or c.startswith("|")
                or c == r"\["
                or c.startswith("![")
            ):
                break
            block.append(c)
            i += 1

        if not block:
            i += 1
            continue

        first_list = re.match(r"^(\d+)\.\s+(.+)$", block[0])
        first_bullet = re.match(r"^[-*+]\s+(.+)$", block[0])
        if first_list or first_bullet:
            ordered = first_list is not None
            item_pattern = re.compile(r"^\d+\.\s+(.+)$" if ordered else r"^[-*+]\s+(.+)$")
            items: list[str] = []
            current: list[str] = []
            for part in block:
                match = item_pattern.match(part)
                if match:
                    if current:
                        items.append(" ".join(current))
                    current = [match.group(1)]
                else:
                    current.append(part)
            if current:
                items.append(" ".join(current))
            tokens.append(Token("list", {"ordered": ordered, "items": items}))
        else:
            hard_break_parts: list[str] = []
            for part in block:
                hard_break_parts.append(part.rstrip())
            tokens.append(Token("paragraph", " ".join(hard_break_parts)))

    return metadata, tokens


INLINE_RE = re.compile(
    r"(\\\(.+?\\\)|\*\*.+?\*\*|`.+?`|(?<!\*)\*[^*]+?\*(?!\*)|\[[^\]]+\]\([^\)]+\))"
)


def add_inline(paragraph, text: str, *, default_size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            if default_size:
                set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith(r"\("):
            # Markdown source uses TeX's inline-math delimiters for a small
            # number of simple identifiers.  Preserve the mathematical text
            # while ensuring the delimiters themselves are not printed.
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, name="Cambria Math", size=default_size, italic=True)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(default_size or 11) - 0.5, color="31465A")
        elif token.startswith("["):
            link = re.fullmatch(r"\[([^\]]+)\]\(([^\)]+)\)", token)
            label = link.group(1) if link else token
            target = link.group(2) if link else ""
            run = paragraph.add_run(label)
            set_run_font(run, size=default_size, color="2E74B5")
            run.underline = True
            if target:
                tail = paragraph.add_run(f" ({target})")
                set_run_font(tail, size=default_size, color="66717E")
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if default_size:
            set_run_font(run, size=default_size)


def add_equation(doc: Document, source: str) -> None:
    paragraph = doc.add_paragraph(style="Equation")
    source = source.replace(r"\,", " ")
    i = 0
    buffer: list[str] = []
    while i < len(source):
        if source[i] == "_":
            if buffer:
                run = paragraph.add_run("".join(buffer))
                set_run_font(run, name="Cambria Math", size=10.5, color=PRESET["colors"]["ink"])
                buffer = []
            i += 1
            if i < len(source) and source[i] == "{":
                end = source.find("}", i + 1)
                if end == -1:
                    sub = source[i + 1 :]
                    i = len(source)
                else:
                    sub = source[i + 1 : end]
                    i = end + 1
            elif i < len(source):
                sub = source[i]
                i += 1
            else:
                sub = ""
            run = paragraph.add_run(sub)
            set_run_font(run, name="Cambria Math", size=8.5, color=PRESET["colors"]["ink"])
            run.font.subscript = True
        else:
            buffer.append(source[i])
            i += 1
    if buffer:
        run = paragraph.add_run("".join(buffer))
        set_run_font(run, name="Cambria Math", size=10.5, color=PRESET["colors"]["ink"])


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    # Field characters and instructions must each live inside a run.  Direct
    # children of w:p are tolerated by some readers but LibreOffice will then
    # freeze the cached display value at "1" on every page.
    run_begin = OxmlElement("w:r")
    run_begin.append(begin)
    run_instr = OxmlElement("w:r")
    run_instr.append(instr)
    run_separate = OxmlElement("w:r")
    run_separate.append(separate)
    run_end = OxmlElement("w:r")
    run_end.append(end)
    paragraph._p.extend([run_begin, run_instr, run_separate, text_run, run_end])


def clear_container(container) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def set_header_footer(section, *, cover: bool = False) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    usable = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(usable, WD_ALIGN_PARAGRAPH.RIGHT)
    if cover:
        left, right = "MARGINAL EDUCATIONAL ACCESS", "RESEARCH REPORT"
    else:
        left, right = "AI COMPUTE FOR EDUCATIONAL ACCESS", "RESEARCH REPORT"
    r1 = hp.add_run(left)
    set_run_font(r1, size=8.5, color=PRESET["colors"]["muted"], bold=True)
    r2 = hp.add_run("\t" + right)
    set_run_font(r2, size=8.5, color=PRESET["colors"]["muted"])

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    if cover:
        run = fp.add_run("OPEN EDUCATIONAL ACCESS RESEARCH")
        set_run_font(run, size=8, color=PRESET["colors"]["muted"])
    else:
        run = fp.add_run("Page ")
        set_run_font(run, size=8.5, color=PRESET["colors"]["muted"])
        add_field(fp, "PAGE")


def configure_section(section, orientation: str, *, cover: bool = False) -> None:
    page = PRESET["page"]
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(page["landscape_width_in"])
        section.page_height = Inches(page["landscape_height_in"])
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(page["portrait_width_in"])
        section.page_height = Inches(page["portrait_height_in"])
    section.top_margin = Inches(page["margin_in"])
    section.bottom_margin = Inches(page["margin_in"])
    section.left_margin = Inches(page["margin_in"])
    section.right_margin = Inches(page["margin_in"])
    section.header_distance = Inches(page["header_footer_in"])
    section.footer_distance = Inches(page["header_footer_in"])
    set_header_footer(section, cover=cover)


def restart_page_number(section, number: int) -> None:
    sectpr = section._sectPr
    pgnum = sectpr.find(qn("w:pgNumType"))
    if pgnum is None:
        pgnum = OxmlElement("w:pgNumType")
        sectpr.append(pgnum)
    pgnum.set(qn("w:start"), str(number))


def clear_page_number_restart(section) -> None:
    """Prevent copied section properties from restarting every section at 1."""
    pgnum = section._sectPr.find(qn("w:pgNumType"))
    if pgnum is not None:
        section._sectPr.remove(pgnum)


def add_cover(doc: Document, metadata: dict[str, str]) -> None:
    section = doc.sections[0]
    configure_section(section, "portrait", cover=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(78)
    spacer.paragraph_format.space_after = Pt(0)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("RESEARCH REPORT")
    set_run_font(run, size=10.5, color=PRESET["colors"]["gold"], bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run(metadata["title"])
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(metadata.get("subtitle", ""))

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(10)
    descriptor.paragraph_format.space_after = Pt(72)
    run = descriptor.add_run("A global portfolio model for language, shared-core, accessibility, and open-curriculum decisions")
    set_run_font(run, size=10.5, color=PRESET["colors"]["gold"], italic=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run(metadata.get("date", ""))
    set_run_font(run, size=12, color=PRESET["colors"]["ink"], bold=True)


def add_content_section(doc: Document, orientation: str) -> object:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    clear_page_number_restart(section)
    configure_section(section, orientation)
    return section


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def keep_table_row_on_one_page(row) -> None:
    """Prevent orphaned cell fragments when a table crosses a page boundary."""
    trpr = row._tr.get_or_add_trPr()
    cant_split = trpr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def is_numeric_cell(text: str) -> bool:
    cleaned = text.strip().replace(",", "").replace("$", "").replace("%", "")
    cleaned = cleaned.replace("–", "-").replace("—", "")
    return bool(cleaned) and bool(re.fullmatch(r"[-+]?\d+(?:\.\d+)?(?:\s*/\s*\d+)*", cleaned))


def compute_column_widths(rows: Sequence[Sequence[str]], total_dxa: int) -> list[int]:
    headers = [h.lower() for h in rows[0]]
    max_lengths = [max(len(str(row[i])) for row in rows) for i in range(len(headers))]
    weights: list[float] = []
    narrow_terms = {"pos.", "position", "rows", "outputs", "depth", "year", "year(s)", "scenario"}
    wide_fragments = ("intervention", "measure", "classification", "product", "evidence", "factor", "artifact", "scope", "task bundle")
    for header, length in zip(headers, max_lengths):
        weight = 4.0 + 2.8 * math.sqrt(max(1, min(length, 120)))
        if header in narrow_terms:
            weight *= 0.48
        if any(fragment in header for fragment in wide_fragments):
            weight *= 1.32
        if "source id" in header or header == "profile" or "profile(s)" in header:
            weight *= 0.72
        weights.append(max(3.0, weight))

    minimum = 520 if len(headers) >= 7 else 700
    widths = [max(minimum, int(total_dxa * weight / sum(weights))) for weight in weights]
    while sum(widths) > total_dxa:
        idx = max(range(len(widths)), key=lambda i: widths[i] - minimum)
        reducible = widths[idx] - minimum
        if reducible <= 0:
            break
        widths[idx] -= min(reducible, sum(widths) - total_dxa)
    if sum(widths) < total_dxa:
        widest = max(range(len(widths)), key=lambda i: weights[i])
        widths[widest] += total_dxa - sum(widths)
    # Resolve any last integer drift deterministically in the final column.
    widths[-1] += total_dxa - sum(widths)
    return widths


def apply_table_geometry(table, widths: Sequence[int], total_dxa: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblpr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar", "w:tblBorders"):
        existing = tblpr.find(qn(tag))
        if existing is not None:
            tblpr.remove(existing)

    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), str(total_dxa))
    tblw.set(qn("w:type"), "dxa")
    tblpr.append(tblw)
    tblind = OxmlElement("w:tblInd")
    tblind.set(qn("w:w"), str(PRESET["tables"]["indent_dxa"]))
    tblind.set(qn("w:type"), "dxa")
    tblpr.append(tblind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)

    cell_mar = OxmlElement("w:tblCellMar")
    for side, value in (
        ("top", PRESET["tables"]["cell_top_bottom_dxa"]),
        ("start", PRESET["tables"]["cell_start_end_dxa"]),
        ("bottom", PRESET["tables"]["cell_top_bottom_dxa"]),
        ("end", PRESET["tables"]["cell_start_end_dxa"]),
    ):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        cell_mar.append(node)
    tblpr.append(cell_mar)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), PRESET["tables"]["border"])
        borders.append(node)
    tblpr.append(borders)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Twips(widths[index])
            set_cell_margins(
                cell,
                PRESET["tables"]["cell_top_bottom_dxa"],
                PRESET["tables"]["cell_start_end_dxa"],
                PRESET["tables"]["cell_top_bottom_dxa"],
                PRESET["tables"]["cell_start_end_dxa"],
            )


def add_table(doc: Document, rows: Sequence[Sequence[str]], *, landscape: bool) -> None:
    column_count = len(rows[0])
    total = PRESET["page"]["landscape_width_dxa"] if landscape else PRESET["page"]["portrait_width_dxa"]
    is_interlanguage_summary = [str(cell).strip() for cell in rows[0]] == [
        "ID",
        "Mechanism",
        "Rows",
        "Exact communities/varieties",
        "Scripts",
        "Task evidence",
        "Prior-study evidence",
        "Existing-edition overlap",
        "Exclusions",
        "Double-count rule",
    ]
    # Appendix H contains one unusually evidence-dense Interslavic row.  At the
    # standard landscape-table size that row is taller than a page, so Word
    # must split it even when w:cantSplit is present.  A narrow, named 7 pt
    # override preserves the full comparison table while keeping every record
    # whole and legible on a landscape page.
    font_size = 7.0 if is_interlanguage_summary else (
        7.5 if landscape and column_count >= 7 else (8.0 if landscape else 8.5)
    )
    widths = compute_column_widths(rows, total)

    table = doc.add_table(rows=len(rows), cols=column_count)
    apply_table_geometry(table, widths, total)
    set_repeat_table_header(table.rows[0])
    for row_index, row_values in enumerate(rows):
        keep_table_row_on_one_page(table.rows[row_index])
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, PRESET["tables"]["header_fill"])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.widow_control = True
            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_numeric_cell(str(value)):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif len(str(value)) <= 16 and column_index != 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, str(value), default_size=font_size)
            for run in paragraph.runs:
                set_run_font(run, size=font_size, bold=(True if row_index == 0 else run.bold))

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def table_is_wide(rows: Sequence[Sequence[str]], context: str) -> bool:
    columns = len(rows[0])
    longest = max(len(str(cell)) for row in rows for cell in row)
    if context in {"appendix_b", "appendix_c"}:
        return True
    if columns >= 6:
        return True
    if columns >= 4 and longest >= 54:
        return True
    if columns == 5 and sum(max(len(row[i]) for row in rows) for i in range(columns)) > 120:
        return True
    return False


def token_contexts(tokens: Sequence[Token]) -> list[str]:
    contexts: list[str] = []
    context = "body"
    for token in tokens:
        if token.kind == "heading":
            lower = str(token.value).lower()
            if lower == "abstract":
                context = "abstract"
            elif lower == "references":
                context = "references"
            elif lower.startswith("appendix a"):
                context = "appendix_a"
            elif lower.startswith("appendix b"):
                context = "appendix_b"
            elif lower.startswith("appendix c"):
                context = "appendix_c"
            elif token.level == 1:
                context = "body"
        contexts.append(context)
    return contexts


def orientation_plan(tokens: Sequence[Token]) -> list[str]:
    """Keep a heading, its explanatory prose, and its wide table together.

    Marking only the table landscape strands a few lines of introductory prose
    on a mostly blank portrait page.  Instead, the nearest enclosing subsection
    becomes a landscape reading unit.  Prose within that unit is narrowed by a
    named gutter override; only tables use the full nine-inch measure.
    """
    contexts = token_contexts(tokens)
    plan = ["portrait"] * len(tokens)
    for index, context in enumerate(contexts):
        if context in {"appendix_b", "appendix_c"}:
            plan[index] = "landscape"

    for table_index, token in enumerate(tokens):
        if token.kind != "table" or not table_is_wide(token.value, contexts[table_index]):
            continue
        heading_index = next(
            (i for i in range(table_index, -1, -1) if tokens[i].kind == "heading"),
            table_index,
        )
        level = tokens[heading_index].level if tokens[heading_index].kind == "heading" else 1
        end = len(tokens)
        for future in range(table_index + 1, len(tokens)):
            if tokens[future].kind == "heading" and tokens[future].level <= level:
                end = future
                break
        for i in range(heading_index, end):
            plan[i] = "landscape"
    # Results, implementation, limitations, and conclusion form one continuous
    # analytic plate. Keeping them landscape prevents a short tail paragraph
    # from being stranded when the report returns to portrait for references.
    for heading_index, token in enumerate(tokens):
        if token.kind != "heading" or token.level != 1:
            continue
        if not re.match(r"^(7\. Results|8\. Grant|9\. Limitations|10\. Conclusion)\b", str(token.value)):
            continue
        end = next(
            (
                i
                for i in range(heading_index + 1, len(tokens))
                if tokens[i].kind == "heading" and tokens[i].level <= 1
            ),
            len(tokens),
        )
        for i in range(heading_index, end):
            plan[i] = "landscape"
    # A parent heading immediately followed by a landscape child heading should
    # travel with that child instead of becoming an orphan on the prior page.
    for index in range(len(tokens) - 1):
        if (
            tokens[index].kind == "heading"
            and tokens[index + 1].kind == "heading"
            and plan[index + 1] == "landscape"
        ):
            plan[index] = "landscape"
    return plan


def apply_landscape_prose_override(paragraph, *, list_item: bool = False) -> None:
    """Restore a readable 6.5-inch prose measure inside a 9-inch section."""
    paragraph.paragraph_format.right_indent = Inches(1.25)
    if list_item:
        paragraph.paragraph_format.left_indent = Inches(1.625)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    else:
        paragraph.paragraph_format.left_indent = Inches(1.25)


def add_image(doc: Document, image_path: Path, alt_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.15))
    docpr = shape._inline.docPr
    docpr.set("descr", alt_text or ALT_TEXT_DEFAULT)
    docpr.set("title", "Figure: rank sensitivity")


def add_body_paragraph(doc: Document, text: str, context: str) -> None:
    plain = re.sub(r"[*`]", "", text).strip()
    is_caption = bool(re.match(r"^(Table\s+\d|Table\s+\d+[A-Z]|Figure\s+\d)", plain, re.I))
    if is_caption:
        paragraph = doc.add_paragraph(style="Caption")
    elif context == "references":
        paragraph = doc.add_paragraph(style="Reference Entry")
    elif plain.lower().startswith(("note.", "source:", "sources:")):
        paragraph = doc.add_paragraph(style="Table Source Note")
    else:
        paragraph = doc.add_paragraph(style="Normal")
    add_inline(paragraph, text)
    paragraph.paragraph_format.widow_control = True
    if plain.lower().startswith("figure "):
        # The preceding image already keeps this caption with itself. Do not
        # also chain the caption to the next table and push all three objects.
        paragraph.paragraph_format.keep_with_next = False
    if context == "abstract" and not plain.lower().startswith("keywords:"):
        # Named override: the shaded abstract is left-aligned to avoid the wide
        # word-space rivers that justified copy can create in a bordered block.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.12)
        set_paragraph_shading(paragraph, PRESET["colors"]["abstract_fill"])
        set_paragraph_border(paragraph, "left", PRESET["colors"]["abstract_border"], size=18)


def build_document(source: Path, figure: Path | None, output: Path) -> dict[str, object]:
    metadata, tokens = parse_markdown(source.read_text(encoding="utf-8-sig"))
    layout_plan = orientation_plan(tokens)
    doc = Document()
    configure_styles(doc)
    numbering = NumberingFactory(doc)

    fixed_time = datetime(2026, 8, 30, 0, 0, tzinfo=timezone.utc)
    props = doc.core_properties
    props.title = metadata["title"]
    props.subject = metadata.get("subtitle", "")
    props.author = "OpenAI Codex gpt-5.6-sol, Ultra"
    props.keywords = "educational access; language; open educational resources; AI compute"
    props.comments = "Built from PAPER.md using the narrative_proposal design preset."
    props.created = fixed_time
    props.modified = fixed_time

    add_cover(doc, metadata)
    current_section = add_content_section(doc, "portrait")
    restart_page_number(current_section, 1)
    orientation = "portrait"
    context = "body"
    table_count = 0
    image_count = 0
    list_count = 0

    for index, token in enumerate(tokens):
        if token.kind == "heading":
            heading_text = str(token.value)
            lower = heading_text.lower()
            if lower == "abstract":
                context = "abstract"
            elif lower == "references":
                context = "references"
            elif lower.startswith("appendix a"):
                context = "appendix_a"
            elif lower.startswith("appendix b"):
                context = "appendix_b"
            elif lower.startswith("appendix c"):
                context = "appendix_c"
            elif token.level == 1:
                context = "body"

        wanted = layout_plan[index]
        section_changed = wanted != orientation
        if section_changed:
            current_section = add_content_section(doc, wanted)
            orientation = wanted

        if token.kind == "heading":
            heading_text = str(token.value)
            paragraph = doc.add_paragraph(style=f"Heading {token.level}")
            if (
                not section_changed
                and (
                    heading_text.lower() in {"references"}
                    or heading_text.lower().startswith("appendix")
                )
            ):
                # Page-break-before is idempotent when the heading already lands
                # at the top of a page.  An explicit empty page-break paragraph
                # can instead create a genuinely blank page when the preceding
                # table ends exactly at a page boundary.
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, heading_text)
            if orientation == "landscape":
                apply_landscape_prose_override(paragraph)
        elif token.kind == "paragraph":
            add_body_paragraph(doc, str(token.value), context)
            if orientation == "landscape":
                apply_landscape_prose_override(doc.paragraphs[-1])
        elif token.kind == "equation":
            add_equation(doc, str(token.value))
            if orientation == "landscape":
                apply_landscape_prose_override(doc.paragraphs[-1])
        elif token.kind == "list":
            ordered = bool(token.value["ordered"])
            num_id = numbering.new_instance(ordered)
            list_count += 1
            for item in token.value["items"]:
                paragraph = doc.add_paragraph(style="Proposal List")
                apply_numbering(paragraph, num_id)
                add_inline(paragraph, str(item))
                if orientation == "landscape":
                    apply_landscape_prose_override(paragraph, list_item=True)
        elif token.kind == "table":
            table_count += 1
            add_table(doc, token.value, landscape=(orientation == "landscape"))
        elif token.kind == "image":
            resolved = figure if figure is not None else (source.parent / token.value["path"])
            if not resolved.exists():
                raise FileNotFoundError(f"Figure not found: {resolved}")
            add_image(doc, resolved, str(token.value.get("alt") or ALT_TEXT_DEFAULT))
            image_count += 1

    # Avoid a floating trailing section-break paragraph where possible.
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return {
        "source": str(source.resolve()),
        "figure": str(figure.resolve()) if figure else None,
        "output": str(output.resolve()),
        "preset": PRESET["name"],
        "header_pattern": "editorial_cover",
        "named_overrides": [
            "landscape_wide_table",
            "appendix_h_dense_comparison_table",
            "abstract_left_rule",
            "reference_hanging_indent",
        ],
        "tokens": len(tokens),
        "tables": table_count,
        "images": image_count,
        "list_blocks": list_count,
        "sections": len(doc.sections),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="PAPER.md input")
    parser.add_argument("--figure", type=Path, help="Optional override for the manuscript figure")
    parser.add_argument("--output", required=True, type=Path, help="DOCX output path")
    parser.add_argument("--report-json", type=Path, help="Optional build summary JSON")
    args = parser.parse_args()

    source = args.input.resolve()
    figure = args.figure.resolve() if args.figure else None
    output = args.output.resolve()
    if source == output:
        raise SystemExit("Input and output must differ")
    report = build_document(source, figure, output)
    if args.report_json:
        args.report_json.parent.mkdir(parents=True, exist_ok=True)
        args.report_json.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
